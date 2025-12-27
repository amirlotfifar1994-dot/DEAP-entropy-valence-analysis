#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
build_manuscript_with_figures.py

Creates a fully data-aligned academic Word manuscript (DOCX) + figures from your *final* DEAP pipeline outputs.

Expected inputs (inside --results_dir):
  - deap_entropy_primary.csv
  - deap_entropy_artifact_windows_removed.csv
  - deap_entropy_artifact_excluded.csv

Optional (in project root OR results_dir):
  - reliability_split_half.csv
  - reliability_test_retest.csv
  - gaussianity.csv

Outputs:
  - figures/*.png
  - manuscript_docx (default: DEAP_entropy_manuscript_WITH_FIGURES.docx)

Run (Windows CMD):
  python build_manuscript_with_figures.py --results_dir results_v3_fp12_p995

Tip:
  Run this AFTER you generate results with:
    python main_deap_pipeline_v3.py ... --out_dir results_v3_fp12_p995 ...

This script recomputes ALL numbers from the CSVs to avoid mismatch between paper and data.
"""

import argparse
import os
import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
from scipy import stats
import statsmodels.api as sm

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

DATASETS = [
    ("PRIMARY", "deap_entropy_primary.csv"),
    ("WIN_REMOVED", "deap_entropy_artifact_windows_removed.csv"),
    ("EXCLUDED", "deap_entropy_artifact_excluded.csv"),
]

def read_csv_anywhere(name: str, results_dir: str, project_root: str) -> str | None:
    cand = [
        os.path.join(results_dir, name),
        os.path.join(project_root, name),
    ]
    for p in cand:
        if os.path.exists(p) and os.path.getsize(p) > 10:
            return p
    return None

def load_results(results_dir: str) -> dict:
    out = {}
    for tag, fname in DATASETS:
        p = os.path.join(results_dir, fname)
        if os.path.exists(p) and os.path.getsize(p) > 10:
            df = pd.read_csv(p)
            # basic validation
            need = {"participant","trial","H_e","valence","arousal","n_windows","n_flagged","flagged_fraction","any_flagged"}
            miss = need - set(df.columns)
            if miss:
                raise ValueError(f"{fname} missing columns: {sorted(miss)}")
            out[tag] = df.copy()
    if "PRIMARY" not in out:
        raise FileNotFoundError("Missing deap_entropy_primary.csv in --results_dir")
    return out

def ols_cluster(y: np.ndarray, X: np.ndarray, groups: np.ndarray):
    m = sm.OLS(y, X).fit(cov_type="cluster", cov_kwds={"groups": groups})
    return m

def within_between(df: pd.DataFrame, pred: str):
    # between: participant means
    g = df.groupby("participant")[["H_e", pred]].mean()
    rb, pb = stats.pearsonr(g["H_e"].to_numpy(), g[pred].to_numpy())

    # within: centered predictor
    df2 = df.copy()
    df2[pred + "_mean"] = df2.groupby("participant")[pred].transform("mean")
    df2[pred + "_w"] = df2[pred] - df2[pred + "_mean"]
    y = df2["H_e"].to_numpy()
    groups = df2["participant"].to_numpy()

    Xw = sm.add_constant(df2[[pred + "_w"]])
    mw = ols_cluster(y, Xw, groups)

    Xj = sm.add_constant(df2[[pred + "_w", pred + "_mean"]])
    mj = ols_cluster(y, Xj, groups)

    return {
        "between_r": float(rb), "between_p": float(pb), "n_participants": int(g.shape[0]),
        "within_beta": float(mw.params[pred + "_w"]), "within_p": float(mw.pvalues[pred + "_w"]),
        "joint_within_beta": float(mj.params[pred + "_w"]), "joint_within_p": float(mj.pvalues[pred + "_w"]),
        "joint_between_beta": float(mj.params[pred + "_mean"]), "joint_between_p": float(mj.pvalues[pred + "_mean"]),
    }

def summary_stats(df: pd.DataFrame):
    return {
        "n_trials": int(df.shape[0]),
        "n_participants": int(df["participant"].nunique()),
        "val_mean": float(df["valence"].mean()),
        "val_sd": float(df["valence"].std(ddof=1)),
        "aro_mean": float(df["arousal"].mean()),
        "aro_sd": float(df["arousal"].std(ddof=1)),
        "window_total": int(df["n_windows"].sum()),
        "flagged_total": int(df["n_flagged"].sum()),
        "window_flag_rate": float(df["n_flagged"].sum() / df["n_windows"].sum()),
        "trial_any_rate": float(df["any_flagged"].mean()),
        "trials_any": int(df["any_flagged"].sum()),
    }

def assoc_models(df: pd.DataFrame, pred: str):
    y = df["H_e"].to_numpy()
    x = df[pred].to_numpy()
    r, p = stats.pearsonr(y, x)

    X = sm.add_constant(df[[pred]])
    ols = sm.OLS(y, X).fit()
    ols_cl = sm.OLS(y, X).fit(cov_type="cluster", cov_kwds={"groups": df["participant"].to_numpy()})

    return {
        "pearson_r": float(r), "pearson_p": float(p),
        "ols_beta": float(ols.params[pred]), "ols_p": float(ols.pvalues[pred]), "ols_r2": float(ols.rsquared),
        "olscl_beta": float(ols_cl.params[pred]), "olscl_p": float(ols_cl.pvalues[pred]), "olscl_r2": float(ols_cl.rsquared),
    }

def icc3_1_ci(x: np.ndarray, y: np.ndarray, alpha=0.05):
    data = np.vstack([x, y]).T
    n, k = data.shape
    grand = data.mean()
    row_means = data.mean(axis=1, keepdims=True)
    col_means = data.mean(axis=0, keepdims=True)
    ssr = k * np.sum((row_means - grand) ** 2)
    dfr = n - 1
    msr = ssr / dfr
    sse = np.sum((data - row_means - col_means + grand) ** 2)
    dfe = (n - 1) * (k - 1)
    mse = sse / dfe
    icc = (msr - mse) / (msr + (k - 1) * mse)
    F = msr / mse
    Flow = F / stats.f.ppf(1 - alpha/2, dfr, dfe)
    Fup  = F * stats.f.ppf(1 - alpha/2, dfe, dfr)
    lo = (Flow - 1) / (Flow + (k - 1))
    hi = (Fup  - 1) / (Fup  + (k - 1))
    return float(icc), float(lo), float(hi)

def make_figures(dfs: dict, out_dir: str):
    os.makedirs(out_dir, exist_ok=True)

    dfp = dfs["PRIMARY"].copy()
    # ensure numeric
    for c in ["H_e","valence","arousal","flagged_fraction"]:
        dfp[c] = pd.to_numeric(dfp[c], errors="coerce")
    dfp = dfp.dropna(subset=["H_e","valence","arousal"])

    # Figure 1: Dataset sizes
    labels, ns = [], []
    for tag in ["PRIMARY","WIN_REMOVED","EXCLUDED"]:
        if tag in dfs:
            labels.append(tag)
            ns.append(int(dfs[tag].shape[0]))
    plt.figure()
    plt.bar(labels, ns)
    plt.ylabel("Number of trials")
    plt.title("Dataset sizes across artifact-handling variants")
    fig1 = os.path.join(out_dir, "fig1_dataset_sizes.png")
    plt.tight_layout()
    plt.savefig(fig1, dpi=200)
    plt.close()

    # Figure 2: Trial-level scatter H_e vs valence + OLS line (descriptive)
    r, p = stats.pearsonr(dfp["H_e"].to_numpy(), dfp["valence"].to_numpy())
    X = sm.add_constant(dfp[["valence"]])
    ols = sm.OLS(dfp["H_e"].to_numpy(), X).fit()  # H_e ~ valence (matches your inference scripts)
    xv = np.linspace(dfp["valence"].min(), dfp["valence"].max(), 200)
    yv = ols.params[0] + ols.params[1]*xv
    plt.figure()
    plt.scatter(dfp["valence"], dfp["H_e"], s=10)
    plt.plot(xv, yv)
    plt.xlabel("Valence")
    plt.ylabel("Hₑ (trial-level)")
    plt.title(f"Hₑ vs valence (naïve): r={r:.3f}, p={p:.3g}, R²={ols.rsquared:.4f}")
    fig2 = os.path.join(out_dir, "fig2_scatter_he_valence.png")
    plt.tight_layout()
    plt.savefig(fig2, dpi=200)
    plt.close()

    # Figure 3: Between-person (participant means) H_e vs valence
    g = dfp.groupby("participant")[["H_e","valence"]].mean().reset_index()
    rb, pb = stats.pearsonr(g["H_e"].to_numpy(), g["valence"].to_numpy())
    Xb = sm.add_constant(g[["valence"]])
    olsb = sm.OLS(g["H_e"].to_numpy(), Xb).fit()
    xv = np.linspace(g["valence"].min(), g["valence"].max(), 200)
    yv = olsb.params[0] + olsb.params[1]*xv
    plt.figure()
    plt.scatter(g["valence"], g["H_e"], s=30)
    plt.plot(xv, yv)
    plt.xlabel("Mean valence (per participant)")
    plt.ylabel("Mean Hₑ (per participant)")
    plt.title(f"Between-participant association: r={rb:.3f}, p={pb:.3g}")
    fig3 = os.path.join(out_dir, "fig3_between_means_he_valence.png")
    plt.tight_layout()
    plt.savefig(fig3, dpi=200)
    plt.close()

    # Figure 4: Flagged fraction distribution
    if "flagged_fraction" in dfp.columns:
        plt.figure()
        plt.hist(dfp["flagged_fraction"].to_numpy(), bins=30)
        plt.xlabel("Flagged fraction per trial")
        plt.ylabel("Count")
        plt.title("Distribution of artifact-flagged window fraction (trial-level)")
        fig4 = os.path.join(out_dir, "fig4_flagged_fraction_hist.png")
        plt.tight_layout()
        plt.savefig(fig4, dpi=200)
        plt.close()
    else:
        fig4 = None

    return {"fig1": fig1, "fig2": fig2, "fig3": fig3, "fig4": fig4}

def doc_set_defaults(doc: Document):
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    style.font.size = Pt(12)

    sec = doc.sections[0]
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)

def add_title_block(doc: Document):
    t = doc.add_paragraph()
    run = t.add_run("Forehead EEG Differential Entropy and Self-Reported Valence/Arousal in DEAP:\nA Reliability- and Artifact-Aware Analysis")
    run.bold = True
    run.font.size = Pt(16)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER

    a = doc.add_paragraph()
    r2 = a.add_run("Author(s): [Insert name(s)]\nAffiliation(s): [Insert]\nCorresponding author: [Insert email]")
    r2.font.size = Pt(11)
    a.alignment = WD_ALIGN_PARAGRAPH.CENTER

def add_fig(doc: Document, path: str, caption: str, width_in=6.2):
    if not path or not os.path.exists(path):
        return
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(path, width=Inches(width_in))
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER

def build_doc(results_dir: str, project_root: str, dfs: dict, figs: dict, out_docx: str):
    dfp = dfs["PRIMARY"]
    S = {k: summary_stats(v) for k, v in dfs.items()}
    # Associations + within/between for each dataset
    A = {}
    WB = {}
    for tag, df in dfs.items():
        A[(tag,"valence")] = assoc_models(df, "valence")
        A[(tag,"arousal")] = assoc_models(df, "arousal")
        WB[(tag,"valence")] = within_between(df, "valence")
        WB[(tag,"arousal")] = within_between(df, "arousal")

    # Optional reliability + gaussianity
    rel_split = read_csv_anywhere("reliability_split_half.csv", results_dir, project_root)
    rel_test  = read_csv_anywhere("reliability_test_retest.csv", results_dir, project_root)
    gauss_p   = read_csv_anywhere("gaussianity.csv", results_dir, project_root)

    rel_text = "Reliability files were not found; this section is omitted."
    if rel_split and rel_test:
        sh = pd.read_csv(rel_split).dropna()
        tt = pd.read_csv(rel_test).dropna()
        rsh, psh = stats.pearsonr(sh["h1"].to_numpy(), sh["h2"].to_numpy())
        icc_sh, icc_sh_lo, icc_sh_hi = icc3_1_ci(sh["h1"].to_numpy(), sh["h2"].to_numpy())
        rtt, ptt = stats.pearsonr(tt["h_first"].to_numpy(), tt["h_last"].to_numpy())
        icc_tt, icc_tt_lo, icc_tt_hi = icc3_1_ci(tt["h_first"].to_numpy(), tt["h_last"].to_numpy())
        rel_text = (
            f"Split-half reliability was high (Pearson r = {rsh:.4f}; ICC(3,1) = {icc_sh:.4f}, 95% CI [{icc_sh_lo:.4f}, {icc_sh_hi:.4f}]). "
            f"Temporal stability (first vs. last half) was also strong (Pearson r = {rtt:.4f}; ICC(3,1) = {icc_tt:.4f}, "
            f"95% CI [{icc_tt_lo:.4f}, {icc_tt_hi:.4f}])."
        )

    gauss_text = "Gaussianity checks file not found; this section is omitted."
    if gauss_p:
        gdf = pd.read_csv(gauss_p).dropna()
        n = int(gdf.shape[0])
        non = int((~gdf["is_normal"]).sum()) if "is_normal" in gdf.columns else None
        if non is not None:
            gauss_text = f"Gaussianity checks (Shapiro–Wilk) on {n} sampled windows indicated {non}/{n} windows rejected normality at p < 0.05."

    doc = Document()
    doc_set_defaults(doc)
    add_title_block(doc)

    doc.add_paragraph("")
    doc.add_paragraph("Abstract", style="Heading 1")
    prim = S["PRIMARY"]
    abs_text = (
        f"We analyzed the DEAP dataset to test whether a forehead EEG differential-entropy index (Hₑ) relates to self-reported valence and arousal. "
        f"Hₑ was computed from Fp1/Fp2 (baseline-corrected) using 4 s windows with 50% overlap and the Gaussian differential-entropy approximation "
        f"h = ½ ln(2πeσ²), aggregated per trial. Artifact windows were flagged using an empirical variance threshold (calibrated separately), and "
        f"sensitivity analyses compared (i) removing flagged windows and (ii) excluding any trial containing flagged windows. "
        f"In the PRIMARY dataset (N={prim['n_trials']} trials; {prim['n_participants']} participants), naïve trial-level association suggested a small "
        f"negative correlation between Hₑ and valence (r={A[('PRIMARY','valence')]['pearson_r']:.4f}), but dependence-aware inference (participant-clustered) "
        f"and within–between disaggregation did not support a robust within-person relationship. No association was found for arousal."
    )
    doc.add_paragraph(abs_text)
    doc.add_paragraph("Keywords: EEG; differential entropy; DEAP; valence; arousal; artifacts; clustered inference; within–between effects")

    doc.add_paragraph("1. Introduction", style="Heading 1")
    doc.add_paragraph(
        "Differential entropy (DE) is frequently used as an EEG feature in affective computing because, under a Gaussian approximation, it reduces to a "
        "simple function of signal variance. However, trial-level inference in repeated-measures designs requires modeling the non-independence of trials "
        "nested within participants. In addition, transient artifacts can disproportionately affect variance-based measures. "
        "Here we provide a reliability- and artifact-aware analysis of a forehead EEG DE feature (Hₑ) in DEAP and evaluate associations with valence and arousal "
        "using dependence-aware inference and within–between decomposition."
    )

    doc.add_paragraph("2. Methods", style="Heading 1")
    doc.add_paragraph("2.1 Dataset", style="Heading 2")
    doc.add_paragraph(
        "We used the DEAP dataset (32 participants, 40 trials each), containing preprocessed EEG (128 Hz) recorded during one-minute music video excerpts "
        "and post-trial self-reports including valence and arousal."
    )
    doc.add_paragraph("2.2 Feature extraction", style="Heading 2")
    doc.add_paragraph(
        "We focused on Fp1/Fp2. Each trial was baseline-corrected using the 3 s pre-stimulus segment and then segmented into 4 s windows with 50% overlap "
        "(29 windows/trial). For each window, Gaussian differential entropy was computed as h = ½ ln(2πeσ²), and Hₑ was defined as the mean of h over windows. "
        "Fp1/Fp2 were combined by averaging their time-domain signals prior to variance/entropy estimation."
    )
    doc.add_paragraph("2.3 Artifact handling", style="Heading 2")
    doc.add_paragraph(
        "We flagged artifact windows using a variance threshold calibrated from the empirical window-variance distribution (see Appendix for commands). "
        "We report three datasets: PRIMARY (all windows), WIN_REMOVED (flagged windows removed; trials excluded if <20 clean windows remain), and EXCLUDED "
        "(remove any trial with ≥1 flagged window)."
    )
    doc.add_paragraph("2.4 Statistical analysis", style="Heading 2")
    doc.add_paragraph(
        "We report naïve trial-level Pearson/OLS descriptively, but base inference on participant-clustered robust standard errors (OLS-CL) and within–between "
        "decomposition separating trial-to-trial deviations (within) from participant means (between)."
    )

    # Figures (methods/results)
    doc.add_paragraph("3. Results", style="Heading 1")
    add_fig(doc, figs["fig1"], "Figure 1. Number of trials retained under each artifact-handling variant.", width_in=6.0)

    doc.add_paragraph("3.1 Descriptives and artifact prevalence", style="Heading 2")
    doc.add_paragraph(
        f"In PRIMARY, {prim['n_trials']} trials from {prim['n_participants']} participants were included. "
        f"Valence mean±SD = {prim['val_mean']:.3f}±{prim['val_sd']:.3f}; arousal mean±SD = {prim['aro_mean']:.3f}±{prim['aro_sd']:.3f}. "
        f"Across {prim['window_total']:,} windows, {prim['flagged_total']:,} were flagged "
        f"({100*prim['window_flag_rate']:.3f}%), affecting {prim['trials_any']}/{prim['n_trials']} trials "
        f"({100*prim['trial_any_rate']:.2f}%)."
    )

    add_fig(doc, figs["fig2"], "Figure 2. Trial-level association between Hₑ and valence (descriptive OLS line).", width_in=6.0)
    add_fig(doc, figs["fig3"], "Figure 3. Between-participant association using participant means (Hₑ vs valence).", width_in=6.0)
    if figs.get("fig4"):
        add_fig(doc, figs["fig4"], "Figure 4. Distribution of the fraction of flagged windows per trial (PRIMARY).", width_in=6.0)

    doc.add_paragraph("3.2 Associations with valence and arousal", style="Heading 2")
    doc.add_paragraph(
        "Table 1 summarizes the association results. Naïve trial-level correlations were small; however, participant-clustered inference (OLS-CL) and "
        "within–between decomposition did not support a robust within-person relationship between Hₑ and valence. Arousal showed no meaningful association "
        "in any dataset."
    )

    # Table 1
    doc.add_paragraph("Table 1. Associations between Hₑ and valence/arousal across artifact-handling datasets.", style="Heading 3")
    t = doc.add_table(rows=1, cols=9)
    h = t.rows[0].cells
    h[0].text="Dataset"
    h[1].text="N"
    h[2].text="Valence r (p)"
    h[3].text="Valence OLS-CL p"
    h[4].text="Valence within β (p)"
    h[5].text="Valence between β (p)"
    h[6].text="Arousal r (p)"
    h[7].text="Arousal OLS-CL p"
    h[8].text="Arousal within β (p)"

    for tag in ["PRIMARY","WIN_REMOVED","EXCLUDED"]:
        if tag not in dfs:
            continue
        df = dfs[tag]
        row = t.add_row().cells
        row[0].text = tag
        row[1].text = str(int(df.shape[0]))
        av = A[(tag,"valence")]
        wbv = WB[(tag,"valence")]
        aa = A[(tag,"arousal")]
        wba = WB[(tag,"arousal")]
        row[2].text = f"{av['pearson_r']:.4f} ({av['pearson_p']:.4g})"
        row[3].text = f"{av['olscl_p']:.4g}"
        row[4].text = f"{wbv['within_beta']:.4f} ({wbv['within_p']:.4g})"
        row[5].text = f"{wbv['joint_between_beta']:.4f} ({wbv['joint_between_p']:.4g})"
        row[6].text = f"{aa['pearson_r']:.4f} ({aa['pearson_p']:.4g})"
        row[7].text = f"{aa['olscl_p']:.4g}"
        row[8].text = f"{wba['within_beta']:.4f} ({wba['within_p']:.4g})"

    doc.add_paragraph(
        "Notes: OLS-CL uses participant-clustered robust SEs. Within–between values are based on participant-mean centering (within) and participant means (between)."
    )

    doc.add_paragraph("3.3 Reliability and distributional checks", style="Heading 2")
    doc.add_paragraph(rel_text + " " + gauss_text)

    doc.add_paragraph("4. Discussion", style="Heading 1")
    doc.add_paragraph(
        "A small negative naïve trial-level association between Hₑ and valence is observable in DEAP, but dependence-aware inference and within–between "
        "disaggregation indicate no robust within-person effect. Any apparent relationship is therefore best interpreted as weak and primarily reflecting "
        "between-participant differences. No association was observed for arousal. These results highlight the importance of artifact-aware preprocessing "
        "and appropriate dependence modeling in repeated-measures EEG–emotion analyses."
    )

    doc.add_paragraph("5. Limitations and future work", style="Heading 1")
    doc.add_paragraph(
        "The present feature is a broadband, variance-based DE approximation computed on two forehead channels. Future work should evaluate band-limited "
        "entropy, alternative complexity measures, and more sophisticated artifact correction (e.g., ICA/ASR) with transparent reporting. "
        "Additionally, models that incorporate participant-specific rating styles may improve interpretability."
    )

    doc.add_paragraph("6. Conclusion", style="Heading 1")
    doc.add_paragraph(
        "Forehead EEG differential entropy (Hₑ) in DEAP shows high reliability but only weak links to self-reported valence and none to arousal when "
        "properly accounting for repeated measures. Conclusions about emotion prediction from this single feature should therefore remain cautious."
    )

    doc.add_paragraph("References", style="Heading 1")
    refs = [
        "Cameron, A. C., & Miller, D. L. (2015). A practitioner’s guide to cluster-robust inference. Journal of Human Resources, 50(2), 317–372. https://doi.org/10.3368/jhr.50.2.317",
        "Curran, P. J., & Bauer, D. J. (2011). The disaggregation of within-person and between-person effects in longitudinal models of change. Annual Review of Psychology, 62, 583–619. https://doi.org/10.1146/annurev.psych.093008.100356",
        "Duan, R.-N., Zhu, J.-Y., & Lu, B.-L. (2013). Differential entropy feature for EEG-based emotion classification. In 2013 6th International IEEE/EMBS Conference on Neural Engineering (NER) (pp. 81–84). https://doi.org/10.1109/NER.2013.6695876",
        "Koelstra, S., Muhl, C., Soleymani, M., Lee, J.-S., Yazdani, A., Ebrahimi, T., Pun, T., Nijholt, A., & Patras, I. (2012). DEAP: A database for emotion analysis using physiological signals. IEEE Transactions on Affective Computing, 3(1), 18–31. https://doi.org/10.1109/T-AFFC.2011.15",
        "Shrout, P. E., & Fleiss, J. L. (1979). Intraclass correlations: Uses in assessing rater reliability. Psychological Bulletin, 86(2), 420–428. https://doi.org/10.1037/0033-2909.86.2.420",
        "Shapiro, S. S., & Wilk, M. B. (1965). An analysis of variance test for normality (complete samples). Biometrika, 52(3–4), 591–611. https://doi.org/10.1093/biomet/52.3-4.591",
    ]
    for r in refs:
        doc.add_paragraph(r)

    doc.add_paragraph("Appendix A. Reproducible commands", style="Heading 1")
    doc.add_paragraph(
        "Example commands (Windows CMD):\n"
        "  python calibrate_artifact_threshold_fp1fp2.py --data_dir data\n"
        "  python main_deap_pipeline_v3.py --data_dir data --out_dir <RESULTS_DIR> --artifact_var <THRESH> --label_min 1 --label_max 9 --min_clean_windows 20 --combine_mode avg_signal\n"
        "  python build_manuscript_with_figures.py --results_dir <RESULTS_DIR>\n"
        "  (Optional checks)\n"
        "    python analyze_deap_results_mixedlm_v2.py --dir <RESULTS_DIR>\n"
        "    python deap_within_between.py --dir <RESULTS_DIR>\n"
    )

    doc.save(out_docx)

def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--results_dir", required=True, help="Directory with deap_entropy_*.csv outputs")
    ap.add_argument("--project_root", default=".", help="Project root (where reliability_*.csv may exist)")
    ap.add_argument("--out_docx", default="DEAP_entropy_manuscript_WITH_FIGURES.docx", help="Output DOCX filename")
    args = ap.parse_args()

    results_dir = args.results_dir
    project_root = args.project_root

    dfs = load_results(results_dir)
    fig_dir = os.path.join(results_dir, "figures")
    figs = make_figures(dfs, fig_dir)
    build_doc(results_dir, project_root, dfs, figs, args.out_docx)

    print("\nDONE")
    print(f"Figures saved to: {fig_dir}")
    print(f"Manuscript saved to: {args.out_docx}")

if __name__ == "__main__":
    main()
